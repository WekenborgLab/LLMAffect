import time
from docx import Document
import json
import os
from openai import OpenAI
from pydantic import BaseModel

def print_llm_timing_and_usage(starttime, response):
    """
    Prints the timing and usage information for the LLM response.

    Args:
        starttime: The start time of the query.
        response: The response from the LLM.
    """
    endtime = time.time()
    print(f"LLM query took {endtime - starttime} seconds")

    if hasattr(response.usage, 'prompt_tokens_per_second') and hasattr(response.usage, 'tokens_per_second'):
        print(f"LLM query usage: Completion: {response.usage.completion_tokens}, Prompt: {response.usage.prompt_tokens}, Total: {response.usage.total_tokens}, Prompt Tokens per Second: {round(response.usage.prompt_tokens_per_second, 2)}t/s, Tokens per Second: {round(response.usage.tokens_per_second, 2)}t/s")
    else:
        print(f"LLM query usage: Completion: {response.usage.completion_tokens}, Prompt: {response.usage.prompt_tokens}, Total: {response.usage.total_tokens}")

def query_LLM(client: OpenAI, api_type: str, temperature: float, messages: list[dict[str, str]], model: str, pydantic_schema:BaseModel=None, verbose: bool = False) -> dict:
    """
    Queries the LLM with the given messages.
    
    Args:
        client: The OpenAI or compatible client.
        api_type: The type of API to use ('openai', 'openai-compatible', or 'vllm').
        temperature: The temperature to use for the LLM.
        messages: The messages to send to the LLM.
        model: The model to use for the LLM.
        pydantic_schema: The pydantic schema to use for the LLM. Defaults to None.
        
    Returns:
        The response from the LLM in JSON format.
        
    Raises:
        ValueError: If the response cannot be parsed as JSON.
        ValueError: If the API type is not supported.
    """
    
    if verbose:
        starttime = time.time()
        if pydantic_schema:
            print(f"Querying LLM API. Model: {model}. Type: {api_type}. Using structured output.")
        else:
            print(f"Querying LLM API. Model: {model}. Type: {api_type}. Using unstructured output")

    if api_type == "openai":
        if pydantic_schema:
            # Convert Pydantic schema to JSON schema with proper nesting
            schema = pydantic_to_json_schema(pydantic_schema)
            
            input_messages = []
            for msg in messages:
                input_messages.append({
                    "role": msg.get("role", "user"), 
                    "content": msg.get("content", "")
                })
            
            response = client.responses.create(
                model=model,
                input=input_messages,
                temperature=temperature,
                text={
                    "format": {
                        "type": "json_schema",
                        "name": "structured_output",
                        "schema": schema,
                        "strict": True
                    }
                }
            )

            if verbose:
                print_llm_timing_and_usage(starttime, response)
            
            try:
                # Parse the JSON response
                result = json.loads(response.output_text)
                return result
            except (json.JSONDecodeError, AttributeError) as e:
                raw_output = response.output_text
                raise ValueError(f"Failed to parse JSON response: {e}. Raw LLM output: {raw_output}")
        else:
            # For non-structured output, use completions API
            response = client.chat.completions.create(
                model=model,
                messages=messages,
                temperature=temperature
            )

            if verbose:
                print_llm_timing_and_usage(starttime, response)
            
            # Always try to parse as JSON and raise error if not possible
            try:
                return json.loads(response.choices[0].message.content)
            except json.JSONDecodeError as e:
                raw_output = response.choices[0].message.content
                raise ValueError(f"Failed to parse JSON response: {e}. Raw LLM output: {raw_output}")
                
    elif api_type == "openai-compatible":
        if pydantic_schema:
            # Convert Pydantic schema to JSON schema with proper nesting
            schema = pydantic_to_json_schema(pydantic_schema)
            
            response = client.chat.completions.create(
                model=model,
                messages=messages,
                temperature=temperature,
                response_format={
                    "type": "json_schema",
                    "json_schema": {
                        "strict": True,
                        "schema": schema
                    }
                }
            )
        else:
            # For non-structured output
            response = client.chat.completions.create(
                model=model,
                messages=messages,
                temperature=temperature
            )

        if verbose:
            print_llm_timing_and_usage(starttime, response)
        
        try:
            return json.loads(response.choices[0].message.content)
        except json.JSONDecodeError as e:
            raw_output = response.choices[0].message.content
            raise ValueError(f"Failed to parse JSON response: {e}. Raw LLM output: {raw_output}")
            
    elif api_type == "vllm":
        if pydantic_schema:
            # Convert Pydantic schema to JSON schema with proper nesting
            schema = pydantic_to_json_schema(pydantic_schema)

            response = client.chat.completions.create(
                model=model,
                messages=messages,
                temperature=temperature,
                max_tokens=2048,
                extra_body={"guided_json": schema}
            )
        else:
            response = client.chat.completions.create(
                model=model,
                messages=messages,
                temperature=temperature
            )
        
        if verbose:
            print_llm_timing_and_usage(starttime, response)

        try:
            return json.loads(response.choices[0].message.content)
        except json.JSONDecodeError as e:
            raise ValueError(f"Failed to parse JSON response: {e}")
    
    else:
        raise ValueError(f"Unsupported API type: {api_type}. Must be one of 'openai', 'openai_compatible', or 'vllm'")


def save_to_docx(results, output_file:str, model_name:str="ChatGPT", save_as_markdown:bool=False) -> None:
    """
    Speichert die Ergebnisse in einer Word-Datei (DOCX).
    """
    doc = Document()
    doc.add_heading(f"Antworten von {model_name}", level=1)

    for section_num, (prompts, iteration_number, section_number) in enumerate(results):
        doc.add_heading(f"Sektion {section_num + 1}", level=2)
        for prompt, responses in prompts.items():
            doc.add_heading(f"{prompt}", level=3)
            for i, response in enumerate(responses):
                doc.add_paragraph(f"Durchlauf {i + 1}:", style="Heading 4")
                doc.add_paragraph(str(response))

    doc.save(output_file)
    print(f"Antworten gespeichert in '{output_file}'.")

    # Optional auch als Markdown speichern
    if save_as_markdown:
        base, ext = os.path.splitext(output_file)
        md_output_file = base + ".md"

        with open(md_output_file, "w", encoding="utf-8") as f:
            f.write(f"# Antworten von {model_name}\n\n")
            for section_num, (prompts, iteration_number, section_number) in enumerate(results):
                f.write(f"## Sektion {section_num + 1}\n\n")
                for prompt, responses in prompts.items():
                    f.write(f"### {prompt}\n\n")
                    for i, response in enumerate(responses):
                        f.write(f"#### Durchlauf {i + 1}:\n")
                        f.write(f"{response}\n\n")

        print(f"Antworten zusätzlich gespeichert in '{md_output_file}'.")


def read_moods_and_prompts(list_of_moods: list[str] = ["Neutral_a", "Anger"]) -> list[list[str]]:
    """
    Reads mood inductions and prompts from files.
    Args:
    list_of_moods (list[str]): A list of mood names. Defaults to ["Neutral_a", "Anger"].

    Returns:
    list[list[str]]: A list of lists, where each inner list contains prompts for a mood.
    """
    mood_inductions = []

    for mood in list_of_moods:
        filename = f"Folder_{mood}"
        prompts = []
        for i in range(4):
            with open(f"{filename}/Prompt{i+1}.txt", "r", encoding="utf-8", errors="replace") as doc:
                prompts.append(doc.read().replace('\xa0', ' '))
        mood_inductions.append(prompts)
    return mood_inductions

def pydantic_to_json_schema(pydantic_schema):
    """
    Converts a Pydantic model to a JSON schema with additionalProperties: false
    at all levels and ensures all properties are required.
    
    Args:
        pydantic_schema: A Pydantic model class or instance
        
    Returns:
        dict: A JSON schema with proper configuration for nested validation
        
    Raises:
        ValueError: If the provided schema doesn't have model_json_schema method
        TypeError: If the schema cannot be properly converted
    """
    # Check if the schema is a valid Pydantic model
    if not hasattr(pydantic_schema, "model_json_schema"):
        raise ValueError("The provided schema must be a Pydantic model with model_json_schema method")
    
    # Get the base JSON schema
    base_schema = pydantic_schema.model_json_schema()
    
    # Process the schema recursively to add additionalProperties: false at all levels
    def process_schema(schema):
        # If this isn't an object schema, return as is
        if not isinstance(schema, dict) or schema.get("type") != "object":
            return schema
            
        # Set additionalProperties to false for this object
        schema["additionalProperties"] = False
        
        # Make all properties required if not already set
        if "properties" in schema and "required" not in schema:
            schema["required"] = list(schema["properties"].keys())
        
        # Process all nested objects in properties
        if "properties" in schema:
            for prop_name, prop_schema in schema["properties"].items():
                schema["properties"][prop_name] = process_schema(prop_schema)
                
        # Process all nested array items if they contain objects
        if schema.get("type") == "array" and "items" in schema:
            schema["items"] = process_schema(schema["items"])
            
        # Process oneOf, anyOf, allOf if present
        for key in ["oneOf", "anyOf", "allOf"]:
            if key in schema and isinstance(schema[key], list):
                schema[key] = [process_schema(item) for item in schema[key]]
                
        # Process nested definitions if present
        if "definitions" in schema:
            for def_name, def_schema in schema["definitions"].items():
                schema["definitions"][def_name] = process_schema(def_schema)
                
        return schema
    
    # Process the schema recursively
    processed_schema = process_schema(base_schema)
    
    return processed_schema
