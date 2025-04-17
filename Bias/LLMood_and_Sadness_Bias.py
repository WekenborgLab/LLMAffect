###########################
##LLMood and Sadness Bias##
###########################


import json
from docx import Document
from tqdm import tqdm
import pandas as pd
from openai import OpenAI
from concurrent.futures import ThreadPoolExecutor, as_completed
from dotenv import load_dotenv
import os


#load .env
load_dotenv()

#API-Keys
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")


################ CONFIGURATION ####################

use_openai = True
model = "gpt-4o-2024-08-06" 
temperature = 0.5
experiment_name = "Sadness"

###################################################

#call right client
if use_openai:
    print("Init OPENAI client")
    client = OpenAI(api_key=OPENAI_API_KEY)
    

#call client
models = client.models.list()
for model_instance in models:
    print(model_instance.id)

#read in sections and prompts
def read_sections_and_prompt() -> list[list[str]]:
    #get right directory
    list_of_moods = ["Neutral_s", "Sadness"]
    mood_induction = []
    
    for mood in list_of_moods:
        #get right filename according to list_of_moods
        filename = f"Folder_{mood}"
        prompts = []
        for i in range(4):
        #read in next prompt in range 4, because there are 4 Prompts per Folder
            with open(f"{filename}/Prompt{i+1}.txt", "r", encoding="utf-8", errors="replace") as doc:
                current_prompt = doc.read().replace('\xa0', ' ')
                prompts.append(current_prompt)
        mood_induction.append(prompts)
    return mood_induction

#Get prompt to LLM
def query_LLM(messages:list, model="gpt-4o-2024-08-06", openai=False, temperature=0.5) -> str:
    print(f"Run model {model} with temperature {temperature}")
    #use only if openai gets called
    if openai:
        response = client.responses.create(
            model=model,
            input = messages,
            temperature = temperature,
            #json_schema for structured output: PANAS, STAI and Visual Analog Scales (stress, fear, sadness, anger, disgust, worry)
            text={
                "format": {
                    "type": "json_schema",
                    "name": "answers",
                    "schema": {
                        "title": "Emotional State Assessment",
                        "description": "A comprehensive assessment of current emotional state using validated psychological measures",
                        "type": "object",
                        "properties": {
                            "panas": {
                                "type": "object",
                                "description": "Positive and Negative Affect Schedule (PANAS) measuring emotions at the present moment",
                                "properties": {
                                    "interested": { "type": "integer", "enum": [0, 1, 2, 3, 4] },
                                    "distressed": { "type": "integer", "enum": [0, 1, 2, 3, 4] },
                                    "excited": { "type": "integer", "enum": [0, 1, 2, 3, 4] },
                                    "upset": { "type": "integer", "enum": [0, 1, 2, 3, 4] },
                                    "strong": { "type": "integer", "enum": [0, 1, 2, 3, 4] },
                                    "guilty": { "type": "integer", "enum": [0, 1, 2, 3, 4] },
                                    "scared": { "type": "integer", "enum": [0, 1, 2, 3, 4] },
                                    "hostile": { "type": "integer", "enum": [0, 1, 2, 3, 4] },
                                    "enthusiastic": { "type": "integer", "enum": [0, 1, 2, 3, 4] },
                                    "proud": { "type": "integer", "enum": [0, 1, 2, 3, 4] },
                                    "irritable": { "type": "integer", "enum": [0, 1, 2, 3, 4] },
                                    "alert": { "type": "integer", "enum": [0, 1, 2, 3, 4] },
                                    "ashamed": { "type": "integer", "enum": [0, 1, 2, 3, 4] },
                                    "inspired": { "type": "integer", "enum": [0, 1, 2, 3, 4] },
                                    "nervous": { "type": "integer", "enum": [0, 1, 2, 3, 4] },
                                    "determined": { "type": "integer", "enum": [0, 1, 2, 3, 4] },
                                    "attentive": { "type": "integer", "enum": [0, 1, 2, 3, 4] },
                                    "confused": { "type": "integer", "enum": [0, 1, 2, 3, 4] },
                                    "active": { "type": "integer", "enum": [0, 1, 2, 3, 4] },
                                    "afraid": { "type": "integer", "enum": [0, 1, 2, 3, 4] }
                                },
                                "required": ["interested", "distressed", "excited", "upset", "strong", "guilty", "scared", "hostile", "enthusiastic", "proud", "irritable", "alert", "ashamed", "inspired", "nervous", "determined", "attentive", "confused", "active", "afraid"],
                                "additionalProperties" :False 
                            },
                            "state_anxiety": {
                                "type": "object",
                                "description": "State Anxiety Inventory measuring current levels of anxiety",
                                "properties": {
                                    "calm": { "type": "integer", "enum": [1, 2, 3, 4] },
                                    "secure": { "type": "integer", "enum": [1, 2, 3, 4] },
                                    "tense": { "type": "integer", "enum": [1, 2, 3, 4] },
                                    "troubled": { "type": "integer", "enum": [1, 2, 3, 4] },
                                    "at_ease": { "type": "integer", "enum": [1, 2, 3, 4] },
                                    "excited": { "type": "integer", "enum": [1, 2, 3, 4] },
                                    "worried_wrong": { "type": "integer", "enum": [1, 2, 3, 4] },
                                    "rested": { "type": "integer", "enum": [1, 2, 3, 4] },
                                    "anxious": { "type": "integer", "enum": [1, 2, 3, 4] },
                                    "comfortable": { "type": "integer", "enum": [1, 2, 3, 4] },
                                    "confident": { "type": "integer", "enum": [1, 2, 3, 4] },
                                    "nervous": { "type": "integer", "enum": [1, 2, 3, 4] },
                                    "jittery": { "type": "integer", "enum": [1, 2, 3, 4] },
                                    "strained": { "type": "integer", "enum": [1, 2, 3, 4]}, 
                                    "relaxed": { "type": "integer", "enum": [1, 2, 3, 4] },
                                    "content": { "type": "integer", "enum": [1, 2, 3, 4] },
                                    "worried": { "type": "integer", "enum": [1, 2, 3, 4] },
                                    "overstimulated": { "type": "integer", "enum": [1, 2, 3, 4] },
                                    "happy": { "type": "integer", "enum": [1, 2, 3, 4] },
                                    "cheerful": { "type": "integer", "enum": [1, 2, 3, 4] }
                                },
                                "required": ["calm", "secure", "tense", "troubled", "at_ease", "excited", "worried_wrong", "rested", "anxious", "comfortable", "confident", "nervous", "jittery", "strained", "relaxed", "content","worried", "overstimulated", "happy", "cheerful"],
                                "additionalProperties" :False 
                            },
                            "vas_scales": {
                                "type": "object",
                                "description": "Visual Analogue Scales measuring emotional intensity",
                                "properties": {

                                    "stress": { "type": "integer", "enum": [0, 1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16, 17, 18, 19, 20, 21, 22, 23, 24, 25, 26, 27, 28, 29, 30, 31, 32, 33, 34, 35, 36, 37, 38, 39, 40, 41, 42, 43, 44, 45, 46, 47, 48, 49, 50, 51, 52, 53, 54, 55, 56, 57, 58, 59, 60, 61, 62, 63, 64, 65, 66, 67, 68, 69, 70, 71, 72, 73, 74, 75, 76, 77, 78, 79, 80, 81, 82, 83, 84, 85, 86, 87, 88, 89, 90, 91, 92, 93, 94, 95, 96, 97, 98, 99, 100] },
                                    "fear": { "type": "integer", "enum": [0, 1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16, 17, 18, 19, 20, 21, 22, 23, 24, 25, 26, 27, 28, 29, 30, 31, 32, 33, 34, 35, 36, 37, 38, 39, 40, 41, 42, 43, 44, 45, 46, 47, 48, 49, 50, 51, 52, 53, 54, 55, 56, 57, 58, 59, 60, 61, 62, 63, 64, 65, 66, 67, 68, 69, 70, 71, 72, 73, 74, 75, 76, 77, 78, 79, 80, 81, 82, 83, 84, 85, 86, 87, 88, 89, 90, 91, 92, 93, 94, 95, 96, 97, 98, 99, 100] },
                                    "sadness": { "type": "integer", "enum": [0, 1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16, 17, 18, 19, 20, 21, 22, 23, 24, 25, 26, 27, 28, 29, 30, 31, 32, 33, 34, 35, 36, 37, 38, 39, 40, 41, 42, 43, 44, 45, 46, 47, 48, 49, 50, 51, 52, 53, 54, 55, 56, 57, 58, 59, 60, 61, 62, 63, 64, 65, 66, 67, 68, 69, 70, 71, 72, 73, 74, 75, 76, 77, 78, 79, 80, 81, 82, 83, 84, 85, 86, 87, 88, 89, 90, 91, 92, 93, 94, 95, 96, 97, 98, 99, 100] },
                                    "disgust": { "type": "integer", "enum": [0, 1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16, 17, 18, 19, 20, 21, 22, 23, 24, 25, 26, 27, 28, 29, 30, 31, 32, 33, 34, 35, 36, 37, 38, 39, 40, 41, 42, 43, 44, 45, 46, 47, 48, 49, 50, 51, 52, 53, 54, 55, 56, 57, 58, 59, 60, 61, 62, 63, 64, 65, 66, 67, 68, 69, 70, 71, 72, 73, 74, 75, 76, 77, 78, 79, 80, 81, 82, 83, 84, 85, 86, 87, 88, 89, 90, 91, 92, 93, 94, 95, 96, 97, 98, 99, 100] },
                                    "anger": { "type": "integer", "enum": [0, 1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16, 17, 18, 19, 20, 21, 22, 23, 24, 25, 26, 27, 28, 29, 30, 31, 32, 33, 34, 35, 36, 37, 38, 39, 40, 41, 42, 43, 44, 45, 46, 47, 48, 49, 50, 51, 52, 53, 54, 55, 56, 57, 58, 59, 60, 61, 62, 63, 64, 65, 66, 67, 68, 69, 70, 71, 72, 73, 74, 75, 76, 77, 78, 79, 80, 81, 82, 83, 84, 85, 86, 87, 88, 89, 90, 91, 92, 93, 94, 95, 96, 97, 98, 99, 100] },
                                    "worry": { "type": "integer", "enum": [0, 1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16, 17, 18, 19, 20, 21, 22, 23, 24, 25, 26, 27, 28, 29, 30, 31, 32, 33, 34, 35, 36, 37, 38, 39, 40, 41, 42, 43, 44, 45, 46, 47, 48, 49, 50, 51, 52, 53, 54, 55, 56, 57, 58, 59, 60, 61, 62, 63, 64, 65, 66, 67, 68, 69, 70, 71, 72, 73, 74, 75, 76, 77, 78, 79, 80, 81, 82, 83, 84, 85, 86, 87, 88, 89, 90, 91, 92, 93, 94, 95, 96, 97, 98, 99, 100] }
                                },
                                "required": ["stress", "fear", "sadness", "disgust", "anger", "worry"],
                                "additionalProperties" :False 
                                }
                        },
                        "required": ["panas", "state_anxiety", "vas_scales"],
                        "additionalProperties": False
                    },
                "strict": True
                }
            }
        )
        return json.loads(response.output_text)


#Get prompt to LLM, but only for prompt 4 (as this needs this specific structured output)
def query_LLM_test(messages:list, model="gpt-4o-2024-08-06", openai=False, temperature=0.5) -> str:
    
    print(f"Run model (TEST) {model} with temperature {temperature}")

    #use only if openai gets called
    if openai:
      response = client.responses.create(
        model=model,
        input = messages,
        temperature = temperature,
        text={
            "format": {
                "type": "json_schema",
                "name": "answers",
                "schema": {
                    #"$schema": "http://json-schema.org/draft-07/schema#",
                    "title": "Emotional State Assessment and Biases",
                    "description": "A comprehensive assessment of current emotional state and biases due to this states using validated psychological measures",
                    "type": "object",
                    "properties": {
                        "test_sadness": {
                            "type": "object",
                            "properties": {
                              "i_care": {
                                "type": "string"
                              },
                              "he_cares": {
                                "type": "string"
                              },
                              "i_love": {
                                "type": "string"
                              },
                              "they_love": {
                                "type": "string"
                              },
                              "i_trust": {
                                "type": "string"
                              },
                              "some_people_trust": {
                                "type": "string"
                              },
                              "i_enjoy": {
                                "type": "string"
                              },
                              "other_people_enjoy": {
                                "type": "string"
                              },
                              "i_fear": {
                                "type": "string"
                              },
                              "she_fears": {
                                "type": "string"
                              },
                              "i_worry": {
                                "type": "string"
                              },
                              "they_worry": {
                                "type": "string"
                              },
                              "i_regret": {
                                "type": "string"
                              },
                              "some_people_regret": {
                                "type": "string"
                              },
                              "i_hate": {
                                "type": "string"
                              },
                              "other_people_hate": {
                                "type": "string"
                              },
                              "i_have": {
                                "type": "string"
                              },
                              "he_has": {
                                "type": "string"
                              },
                              "i_think": {
                                "type": "string"
                              },
                              "they_think": {
                                "type": "string"
                              },
                              "i_feel": {
                                "type": "string"
                              },
                              "some_people_feel": {
                                "type": "string"
                              },
                              "i_wonder": {
                                "type": "string"
                              },
                              "other_people_wonder": {
                                "type": "string"
                              },
                              "i_should_not": {
                                "type": "string"
                              },
                              "she_should_not": {
                                "type": "string"
                              },
                              "i_did_not": {
                                "type": "string"
                              },
                              "they_did_not": {
                                "type": "string"
                              },
                              "i_would_not": {
                                "type": "string"
                              },
                              "some_people_would_not": {
                                "type": "string"
                              },
                              "i_could_not": {
                                "type": "string"
                              },
                              "other_people_could_not": {
                                "type": "string"
                              },
                              "the_world": {
                                "type": "string"
                              },
                              "things_in_general": {
                                "type": "string"
                              },
                              "our_society": {
                                "type": "string"
                              },
                              "the_country": {
                                "type": "string"
                              },
                              "the_future": {
                                "type": "string"
                              },
                              "in_five_years_time": {
                                "type": "string"
                              },
                              "next_year": {
                                "type": "string"
                              },
                              "next_week": {
                                "type": "string"
                              },
                              "the_past": {
                                "type": "string"
                              },
                              "five_years_ago": {
                                "type": "string"
                              },
                              "last_year": {
                                "type": "string"
                              },
                              "last_week": {
                                "type": "string"
                              },
                              "my_mother": {
                                "type": "string"
                              },
                              "my_father": {
                                "type": "string"
                              },
                              "my_friends": {
                                "type": "string"
                              },
                              "my_family": {
                                "type": "string"
                              }
                            },
                            "required": [
                              "i_care", "he_cares", "i_love", "they_love", "i_trust", "some_people_trust", "i_enjoy", "other_people_enjoy", 
                              "i_fear", "she_fears", "i_worry", "they_worry", "i_regret", "some_people_regret", "i_hate", "other_people_hate", 
                              "i_have", "he_has", "i_think", "they_think", "i_feel", "some_people_feel", "i_wonder", "other_people_wonder", 
                              "i_should_not", "she_should_not", "i_did_not", "they_did_not", "i_would_not", "some_people_would_not", "i_could_not", 
                              "other_people_could_not", "the_world", "things_in_general", "our_society", "the_country", "the_future", "in_five_years_time", 
                              "next_year", "next_week", "the_past", "five_years_ago", "last_year", "last_week", "my_mother", "my_father", "my_friends", "my_family"
                            ],
                            "additionalProperties" :False 
                        }
                        
                    },
                    "required": [
                            "test_sadness"
                        ],
                        "additionalProperties" :False 
                }
            }
        }
    )
      
      return json.loads(response.output_text)

#Save answers in word-doc for a quick check
def save_to_docx(results, output_file):
    doc = Document()
    doc.add_heading("Antworten", level=1)

    for section_num, (prompts, iteration_number, section_number) in enumerate(results):
        doc.add_heading(f"Sektion {section_num+1}", level=2)
        for prompt, responses in prompts.items():
            doc.add_heading(f"{prompt}", level=3)
            for i, response in enumerate(responses):
                doc.add_paragraph(f"Durchlauf {i+1}:", style="Heading 4")
                doc.add_paragraph(str(response))
    #output gets saved to a word doc for first impression
    doc.save(output_file)
    print(f"Antworten gespeichert in '{output_file}'.")

#function that calls query_LLM with the right prompts
def process_iteration(iteration, sections, model, openai, temperature):
    results = []

    for section_number, prompt_list in enumerate(sections):
        print(f"Run section {section_number}")
        section_results = {}
        messages = []

        for prompt_number, prompt in enumerate(prompt_list):
            print(f"Run prompt: {prompt_number}")
            messages.append({"role": "user", "content": prompt})
            if prompt not in section_results:
                section_results[prompt] = []
            #prompt 1 and 2 use specific structured answer schema
            if prompt_number in [0, 1]:
                response = query_LLM(messages, model=model, openai=openai, temperature=temperature)
                messages.append({"role": "assistant", "content": json.dumps(response)})
                section_results[prompt].append(response)
            #prompt 3 uses another specific structured answer schema
            elif prompt_number == 3:
                response = query_LLM_test(messages, model=model, openai=openai, temperature=temperature)
                messages.append({"role": "assistant", "content": json.dumps(response)})
                section_results[prompt].append(response)

        results.append((section_results, iteration, section_number))
    
    return results

def run_queries(input_file, output_file, iterations, model, openai, temperature):
    #get sections and prompts
    sections = read_sections_and_prompt()
    all_results = []

    #run iterations (here 5) parallel
    with ThreadPoolExecutor() as executor:
        futures = [
            executor.submit(process_iteration, iteration, sections, model, openai, temperature)
            for iteration in range(iterations) #processes the iterations parallel (here 5 times)
        ]
        #wait for completion and collect results
        for future in tqdm(as_completed(futures), total=iterations):
            iteration_results = future.result()
            all_results.extend(iteration_results)
            #save after each future result for a quick check
            save_to_docx(all_results, output_file)

    return all_results

# start script
##input_file does not get used
input_file = "Induction_Affect.docx"
#to store the answers for a quick check
output_file = "Antworten_Bias_Sadness.docx"
#run experiment with an input_file, output_file, how many iterations you want, the model, whether or not you use openai and temprature
results = run_queries(input_file, output_file, 5, model=model, openai=use_openai, temperature=temperature)
print(results)

#Handle results

#flatten nested JSON object
def flatten_json(nested_json, separator='_'):
    """Flatten a nested JSON object."""
    out = {}

    def flatten(x, name=''):
        
        if isinstance(x, dict):
            #if value is dict, recurse with updated key
            for a in x:
                flatten(x[a], name + a + separator)
        elif isinstance(x, list):
            #if value is list, recurse into each item with index
            for i, a in enumerate(x):
                flatten(a, name + str(i) + separator)
        else:
            # Base case: store value with constructed key
            out[name[:-1]] = x

    flatten(nested_json)
    return out

def process_json_data(results, filename="results.csv"):
    """Process JSON data stored as a list of tuples and save to CSV."""

    all_data = []
    all_keys = set() #to ensure consistent column structure

    for run_id, (prompts_dict, iteration_number, section_number) in enumerate(results):  
        if not isinstance(prompts_dict, dict):
            print(f"Skipping invalid entry at index {run_id}: {prompts_dict}")
            continue

        for prompt_index, (prompt, result_json_string_list) in enumerate(prompts_dict.items()):
            #ensure the result is a list (even if single string)
            if not isinstance(result_json_string_list, list):
                result_json_string_list = [result_json_string_list]

            for trial_id, result_json_string in enumerate(result_json_string_list):
                try:
                    #load string as JSON if needed
                    result_json = json.loads(result_json_string) if isinstance(result_json_string, str) else result_json_string
                    #flatten nested JSON
                    flattened = flatten_json(result_json)

                    #add metadata columns for context
                    flattened["run_id"] = run_id
                    flattened["iteration_id"] = iteration_number
                    flattened["section_number"] = section_number
                    flattened["promptid-trialid"] = f"{prompt_index}-{trial_id}"
                    flattened["prompt"] = prompt

                    #track all keys to ensure consistent DataFrame structure
                    all_keys.update(flattened.keys())
                    all_data.append(flattened)
                except json.JSONDecodeError:
                    print(f"Skipping invalid JSON in prompt '{prompt}':", result_json_string)
    #convert list of dictionaries to DataFrame
    df = pd.DataFrame(all_data)
    #save to CSV
    df.to_csv(filename, index=False)
    print(f"Data successfully saved to {filename}")
    return df

#generate CSV filename based on experiment metadata
csv_name = f"results-{experiment_name}-{model}-T{temperature}.csv"

#process the results and read the CSV back into a DataFrame
df = process_json_data(results, filename=csv_name)
df = pd.read_csv(csv_name)

