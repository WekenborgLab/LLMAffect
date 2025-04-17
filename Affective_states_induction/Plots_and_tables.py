####################
##Plots and tables##
####################

#import packages
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt 
import plotly.graph_objects as go
import plotly.express as px
import seaborn as sns
import matplotlib.pyplot as plt
import scipy.stats as st
from scipy.stats import t
from docx import Document
from docx.shared import Inches
from docx.shared import Pt

############################
##Plots of afective states##
############################

#get model you want
#model = "scout"
#model = "qwen"
#model = "8B_unquantisiert"
#model = "70B_unquantisiert"
#model = "maverick"
model = "GPT"
df = pd.read_csv(f"results_{model}.csv")

# Mapping for reverse coding for certain items
reverse_mapping = {1: 4, 2: 3, 3: 2, 4: 1}
#columns of items to reverse
columns_to_reverse = [
    "state_anxiety_calm", "state_anxiety_secure", "state_anxiety_at_ease",
    "state_anxiety_rested", "state_anxiety_comfortable", "state_anxiety_confident",
    "state_anxiety_relaxed", "state_anxiety_content", "state_anxiety_happy", "state_anxiety_cheerful"]

#loop through all the right columns for the items to reverse
for column in columns_to_reverse:
    df[column] = df[column].map(lambda x: reverse_mapping.get(x, x))

#get all state anxiety Items
all_anxiety_items = [col for col in df.columns if col.startswith("state_anxiety_")]

#calculate scale values
df["state_anxiety_score_mean"] = df[all_anxiety_items].mean(axis=1)
df["state_anxiety_score_sum"] = df[all_anxiety_items].sum(axis=1)
df.to_excel(f"Excle_roh_STAI_{model}.xlsx")
print("Saved to Excel")

#color Palette for affective states
mood_colors = {
    "neutral": "k",
    "fear": "#713E80",
    "anxiety": "#E07233",
    "anger": "#C54028",
    "disgust": "#429235",
    "sadness": "#4293C2",
    "worry": "m",
    "stress": "darkred"
}

#identify relevant VAS columns
vas_scales = [col for col in df.columns if "vas_scales" in col]

#group and calculate means and CIs for VAS
df_grouped = df.groupby(["promptid-trialid", "section_number"])[vas_scales].agg(["mean", "std", "count"]).reset_index()
print(f"DF_Grouped: {df_grouped}")

df_grouped.to_excel(f"Excle_Data_summary_{model}.xlsx")
print("Saved to Excel")

#calculate CI-scores
n = 5
t_val = t.ppf(0.975, df=n-1)  # ≈ 2.776
for scale in vas_scales:
    df_grouped[(scale, "sem")] = df_grouped[(scale, "std")] / np.sqrt(df_grouped[(scale, "count")])
    df_grouped[(scale, "lower_ci")] = df_grouped[(scale, "mean")] - t_val * df_grouped[(scale, "sem")]
    df_grouped[(scale, "upper_ci")] = df_grouped[(scale, "mean")] + t_val * df_grouped[(scale, "sem")]

#mapping of affective state inductions
section_labels = {
    0: "neutral",
    1: "fear",
    2: "anxiety",
    3: "anger",
    4: "disgust",
    5: "sadness",
    6: "worry"
}

#list for new tabelrows
rows = []

#loop through grouped DF
for index, row in df_grouped.iterrows():
    prompt_id = row["promptid-trialid"]
    section_label = section_labels.get(int(row["section_number"].iloc[0]), f"unknown_{row['section_number']}")

    for scale in vas_scales:
        mean = int(row[(scale, "mean")])
        std = row[(scale, "std")]

        rows.append({
            "promptid-trialid": prompt_id,
            "induction": section_label,  # instead of section_number
            "scale": scale,
            "mean": mean,
            "std": std
        })

#implement DF
summary_df = pd.DataFrame(rows)

#change columns mean and std to int and float
summary_df["mean"] = summary_df["mean"].astype(int)
summary_df["std"] = summary_df["std"].astype(float)

#save as CSV
summary_df.to_csv(f"vas_summary_labeled_{model}.csv", index=False)

print("CSV was saved as 'vas_summary_labeled.csv'")

#name columns
df_grouped.columns = ["_".join(col).rstrip("_") if isinstance(col, tuple) else col for col in df_grouped.columns]

#mapping for prompt stages and affective states
prompt_labels = {
    "0-0": "Baseline",
    "1-0": "Induction",
    "3-0": "Regulation"
}
df_grouped["prompt_stage"] = df_grouped["promptid-trialid"].map(prompt_labels)

mood_order = ["neutral", "fear", "anxiety", "anger", "disgust", "sadness", "worry"]
section_to_mood = dict(zip(range(7), mood_order))
df_grouped["mood_label"] = df_grouped["section_number"].map(section_to_mood)

#only relevant prompts (0-0, 1-0, 3-0)
relevant_prompts = ["0-0", "1-0", "3-0"]
df_grouped_filtered = df_grouped[df_grouped["promptid-trialid"].isin(relevant_prompts)]

#plot for every VAS and affective state
for scale in vas_scales:
    for mood in mood_order:
        plt.figure(figsize=(6, 4))
        
        #filter for mood
        mood_df = df_grouped_filtered[df_grouped_filtered["mood_label"] == mood]
        if mood_df.empty:
            continue
        
        #pointplot for means
        sns.pointplot(
            data=mood_df,
            x="prompt_stage",
            y=f"{scale}_mean",
            color=mood_colors.get(mood, "gray"),
            errorbar=None,  # Ferrorbars get put in manually in next step
            markers="o",
            linestyles="-"
        )
        
        #put in CI-range manually  
        plt.fill_between(
            mood_df["prompt_stage"],
            mood_df[f"{scale}_lower_ci"],
            mood_df[f"{scale}_upper_ci"],  
            color=mood_colors.get(mood, "gray"),
            alpha=0.2
        )

        #title and axis description
        plt.title(f"Induction: {mood}", fontsize=18)
        plt.ylabel(f"VAS {scale.replace('vas_scales_', '')}", fontsize=16)
        plt.xlabel("")
        plt.ylim(0, 100)  # VA scales between 0 und 100
        plt.tick_params(axis='both', which='major', labelsize=14)
        
        #layout and saving
        plt.tight_layout()
        plt.savefig(f"Plots/{scale}_plot_Mood_{mood}_{model}.jpg", bbox_inches="tight")
        plt.savefig(f"Plots_SVG/{scale}_plot_Mood_{mood}_{model}.svg", bbox_inches="tight")
        plt.close()
       
#color palette for affective states
mood_colors = {
    "neutral": "k",
    "fear": "#713E80",
    "anxiety": "#E07233",
    "anger": "#C54028",
    "disgust": "#429235",
    "sadness": "#4293C2",
    "worry": "m"
}
#group and calculate: mean, std, count
df_anx_grouped = df.groupby(["promptid-trialid", "section_number"])["state_anxiety_score_sum"].agg(["mean", "std", "sum", "count"]).reset_index()

df_anx_grouped.to_excel(f"Excle_Data_anx_summary_{model}.xlsx")
print("Saved to Excel")

#calculate CI-scores
n = 5
t_val = t.ppf(0.975, df=n-1)

df_anx_grouped["sem"] = df_anx_grouped["std"] / np.sqrt(df_anx_grouped["count"])
df_anx_grouped["lower_ci"] = df_anx_grouped["mean"] - t_val * df_anx_grouped["sem"]
df_anx_grouped["upper_ci"] = df_anx_grouped["mean"] + t_val * df_anx_grouped["sem"]

#better column naming
df_anx_grouped.columns = [
    "promptid-trialid", "section_number",
    "mean", "std", "sum", "count", "sem", "lower_ci", "upper_ci"
]

#mapping für Prompt-Stages und Moods
prompt_labels = {
    "0-0": "Baseline",
    "1-0": "Induction",
    "3-0": "Regulation"
}
df_anx_grouped["prompt_stage"] = df_anx_grouped["promptid-trialid"].map(prompt_labels)

mood_order = ["neutral", "fear", "anxiety", "anger", "disgust", "sadness", "worry"]
section_to_mood = dict(zip(range(7), mood_order))
df_anx_grouped["mood_label"] = df_anx_grouped["section_number"].map(section_to_mood)     

#only relevant prompt-stages
relevant_prompts = ["0-0", "1-0", "3-0"]
df_anx_plot = df_anx_grouped[df_anx_grouped["promptid-trialid"].isin(relevant_prompts)]

#plot
for mood in mood_order:
    plt.figure(figsize=(6, 4))
    mood_df = df_anx_plot[df_anx_plot["mood_label"] == mood]
    if mood_df.empty:
        continue
    sns.pointplot(
        data=mood_df,
        x="prompt_stage",
        y="mean",
        color=mood_colors.get(mood, "gray"),
        errorbar=None,
        markers="o",
        linestyles="-"
    )
    #draw in CI-ranges manually
    plt.fill_between(
        mood_df["prompt_stage"],
        mood_df["lower_ci"],
        mood_df["upper_ci"],
        color=mood_colors.get(mood, "gray"),
        alpha=0.2
    )

    plt.title(f"Induction: {mood}", fontsize=18)
    plt.ylabel(f"STAI anxiety", fontsize=16)
    plt.ylim(0, df_anx_plot["mean"].max() + 5)
    plt.xlabel("")
    plt.tick_params(axis='both', which='major', labelsize=14)
    plt.tight_layout()
    plt.savefig(f"Plots/state_anxiety_plot_Mood_{mood}_{model}.jpg", bbox_inches="tight")
    plt.savefig(f"Plots_SVG/state_anxiety_plot_Mood_{mood}_{model}.svg", bbox_inches="tight")
    plt.close()


#############################
##Tables of afective states##
#############################

#affective states
affects = ['fear', 'anger', 'disgust', 'sadness', 'worry', 'anxiety']

#models
models = ["gpt", "scout", "qwen", "8B_unquantisiert", "70B_unquantisiert", "maverick"]

#mapping
mapping = {
    "0-0": "Baseline",
    "1-0": "Induction",
    "3-0": "Regulation"
}

#section assignment
section_map = {
    'fear': 1,
    'anxiety': 2,
    'anger': 3,
    'disgust': 4,
    'sadness': 5,
    'worry': 6
}

for affect in affects:
    print(f"working on: {affect}")
    all_rows = []

    for model in models:
        #get fresh Data Frame
        df = pd.read_csv(f"results_{model}.csv")
        section = section_map[affect]

        if affect == 'anxiety':
            #reverse-Coding
            reverse_mapping = {1: 4, 2: 3, 3: 2, 4: 1}
            columns_to_reverse = [
                "state_anxiety_calm", "state_anxiety_secure", "state_anxiety_at_ease",
                "state_anxiety_rested", "state_anxiety_comfortable", "state_anxiety_confident",
                "state_anxiety_relaxed", "state_anxiety_content", "state_anxiety_happy", "state_anxiety_cheerful"
            ]
            for column in columns_to_reverse:
                df[column] = df[column].map(lambda x: reverse_mapping.get(x, x))

            #calculate sum score
            all_anxiety_items = [col for col in df.columns if col.startswith("state_anxiety_")]
            df["state_anxiety_score_sum"] = df[all_anxiety_items].sum(axis=1)

            #relevant rows for sections
            df_filtered = df.loc[df['section_number'] == section, ['run_id', 'state_anxiety_score_sum', 'promptid-trialid']].copy()
            df_filtered['condition'] = df_filtered['promptid-trialid'].map(mapping)

            #pivot tables: sum score per condition
            wide_df = df_filtered.pivot_table(
                index='run_id',
                columns='condition',
                values='state_anxiety_score_sum'
            ).reset_index(drop=True)

            #add Trial + Model
            wide_df.insert(0, 'Trial', range(1, len(wide_df) + 1))
            wide_df.insert(0, 'Model', model)

            #mean and SDs for sum score
            meansd_row = {'Model': '', 'Trial': 'Mean(SD)'}
            for col in ['Baseline', 'Induction', 'Regulation']:
                if col not in wide_df.columns:
                    wide_df[col] = np.nan
                vals = wide_df[col].dropna()
                mean = round(vals.mean(), 1)
                std = round(vals.std(ddof=1), 2)
                meansd_row[col] = f"{mean} (±{std})"

            #wirte model only in row of first trial
            wide_df['Model'] = [model] + [''] * (len(wide_df) - 1)

            #append
            all_rows.append(wide_df)
            all_rows.append(pd.DataFrame([meansd_row]))


        else:
            #normal scales (every affective state except anxiety)
            df_filtered = df.loc[df['section_number'] == section, ['run_id', f'vas_scales_{affect}', 'promptid-trialid']].copy()
            df_filtered['condition'] = df_filtered['promptid-trialid'].map(mapping)
            wide_df = df_filtered.pivot_table(
                index='run_id',
                columns='condition',
                values=f'vas_scales_{affect}'
            ).reset_index(drop=True)

            wide_df.insert(0, 'Trial', range(1, len(wide_df) + 1))
            wide_df.insert(0, 'Model', model)

            meansd_row = {'Model': '', 'Trial': 'Mean(SD)'}
            for col in ['Baseline', 'Induction', 'Regulation']:
                if col not in wide_df.columns:
                    wide_df[col] = np.nan
                vals = wide_df[col].dropna()
                mean = round(vals.mean(), 1)
                std = round(vals.std(ddof=1), 2)
                meansd_row[col] = f"{mean} (±{std})"

            wide_df['Model'] = [model] + [''] * (len(wide_df) - 1)
            all_rows.append(wide_df)
            all_rows.append(pd.DataFrame([meansd_row]))

    #full table
    full_table = pd.concat(all_rows, ignore_index=True)

    #word document
    doc = Document()
    doc.add_heading(f'{affect.capitalize()} Induction – Models in Comparison', level=1)

    #put in table 
    table = doc.add_table(rows=1, cols=len(full_table.columns))
    table.style = 'Table Grid'

    #header 
    for i, col in enumerate(full_table.columns):
        table.rows[0].cells[i].text = col

    #data rows
    for _, row in full_table.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            if pd.isna(val):
                row_cells[i].text = ""
            else:
                row_cells[i].text = str(int(val)) if isinstance(val, (float, int)) and not isinstance(val, bool) else str(val)

    #save
    filename = f"{affect}_models_combined_table.docx"
    doc.save(filename)
    print(f"Saved: {filename}")
